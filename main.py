import asyncio
import io
import re
import uuid
from typing import Annotated
from fastapi import FastAPI, HTTPException, status, Header, Query, Request
from fastapi.exceptions import RequestValidationError
from fastapi.encoders import jsonable_encoder
from fastapi.responses import JSONResponse
from docx import Document
from docx.shared import Pt
from datetime import datetime
from docx.enum.text import WD_COLOR_INDEX

from s3_client import s3_client
from config import settings

ru_months = {
    "January" : "января",
    "February" : "февраля",
    "March" : "марта",
    "April" : "апреля",
    "May" : "мая",
    "June" : "июня",
    "July" : "июля",
    "August" : "августа",
    "September" : "сентября",
    "October" : "октября",
    "November" : "ноября",
    "December" : "декабря"
}

def normalize_text(p):
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def replace_placeholders(p, replacements, pattern):
    """Replace every placeholder found in the paragraph in a single pass and
    re-apply the base font. Single-pass so an inserted value that happens to
    contain another placeholder name is not re-substituted."""
    if not pattern.search(p.text):
        return
    p.text = pattern.sub(lambda m: replacements[m.group(0)], p.text)
    normalize_text(p)

app = FastAPI()
items = []
items_id = 1


@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    # Missing required parameters -> single Russian message; the request is
    # rejected before the route runs, so no document is generated.
    if any(error["type"] == "missing" for error in exc.errors()):
        return JSONResponse(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            content={"detail": "Недостаточно необходимых параметров"},
        )
    return JSONResponse(
        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
        content={"detail": jsonable_encoder(exc.errors())},
    )

@app.get('/')
def root():
    return {'message': 'Hello World!'}

@app.post('/generate_docs/')
async def generate_docs(
        customer_inn: Annotated[str, Query(pattern=r'^(\d{10}|\d{12})$')],
        customer_ogrn: Annotated[str, Query(pattern=r'^(\d{13}|\d{15})$')],
        customer_org_name: Annotated[str, Query(min_length=1, max_length=300)],
        tour_date: datetime,
        contract_date: datetime,
        customer_representative: Annotated[str, Query(min_length=1, max_length=300)],
        number_of_visitors: Annotated[int, Query(gt=0)],
        kpp: Annotated[str | None, Query(pattern=r'^\d{9}$')] = None,
        representative_title: Annotated[str, Query(min_length=1, max_length=300)] = 'Генеральный директор',
        x_api_key: str = Header(...)):
    if x_api_key != settings.api_key:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail='Invalid API key',
        )

    # Heavy work (docx generation + blocking S3 upload) runs in a worker
    # thread so it doesn't block the async event loop.
    return await asyncio.to_thread(
        _build_and_upload,
        customer_inn=customer_inn,
        customer_ogrn=customer_ogrn,
        customer_org_name=customer_org_name,
        kpp=kpp,
        tour_date=tour_date,
        contract_date=contract_date,
        customer_representative=customer_representative,
        number_of_visitors=number_of_visitors,
        representative_title=representative_title,
    )


def _build_and_upload(
        customer_inn: str,
        customer_ogrn: str,
        customer_org_name: str,
        kpp: str | None,
        tour_date: datetime,
        contract_date: datetime,
        customer_representative: str,
        number_of_visitors: int,
        representative_title: str) -> dict:
    template = Document('contract_template.docx')
    unique_id = uuid.uuid4().int
    short_uniue_id = str(unique_id)[:8]


    #Date parsing
    contract_day = contract_date.day
    contract_month_ru = ru_months[contract_date.strftime("%B")]
    contract_year = contract_date.year

    tour_day = tour_date.day
    tour_month_ru = ru_months[tour_date.strftime("%B")]
    tour_year = tour_date.year

    replacements = {
        'contract_date': f'«{contract_day}» {contract_month_ru} {contract_year}г.',
        'contract_uuid': short_uniue_id,
        'tour_date': f'«{tour_day}» {tour_month_ru} {tour_year}г',
        'customer_org_name': customer_org_name,
        'customer_representative': customer_representative,
        'representative_title': representative_title,
        'number_of_visitors': str(number_of_visitors),
        'customer_inn': f'ИНН {customer_inn}',
        'customer_ogrn': f'ОГРН {customer_ogrn}',
        'kpp': kpp if kpp else '',
    }
    # Longest token first so a token can't be matched as a prefix of another.
    pattern = re.compile('|'.join(re.escape(k) for k in sorted(replacements, key=len, reverse=True)))

    # Replace placeholders in every top-level paragraph and every table cell
    # so no placeholder token can be left behind in the result.
    for p in template.paragraphs:
        replace_placeholders(p, replacements, pattern)
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders(p, replacements, pattern)
    buffer = io.BytesIO()
    template.save(buffer)
    result = s3_client.upload_file(
        file_content=buffer.getvalue(),
        filename=f"{unique_id}_contract.docx",
        folder="test_folder"
    )
    return result
